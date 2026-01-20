import streamlit as st
import google.generativeai as genai
import json
import os
import base64
import re
from dotenv import load_dotenv
import io
# import ให้ครบ
from docx import Document
from docx.shared import Pt, Inches, RGBColor 
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. โหลด API Key
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

# ตั้งค่าหน้าเว็บ
st.set_page_config(
    layout="wide", 
    page_title="AI Business Planner - DSD Sakon Nakhon",
    page_icon="🛠️"
)

# --- ฟังก์ชันแปลงรูปภาพ ---
def get_img_as_base64(file_path):
    if not os.path.exists(file_path): return ""
    with open(file_path, "rb") as f: data = f.read()
    return base64.b64encode(data).decode()

logo_path = "static/logo_dsd.png"
img_base64 = get_img_as_base64(logo_path)
logo_src = f"data:image/png;base64,{img_base64}" if img_base64 else "https://via.placeholder.com/150?text=Logo"

# --- ข้อมูลตัวอย่าง ---
EXAMPLES_DATA = {
    "🔌 ช่างแอร์/ไฟฟ้า": {
        "name": "ร้านช่างแอร์และไฟฟ้าบริการ",
        "product": "บริการล้างแอร์ ซ่อมแอร์ ติดตั้งระบบไฟ เดินสายไฟ",
        "customer": "เจ้าของบ้านในหมู่บ้านจัดสรร, หอพัก",
        "usp": "ช่างมาไว ไม่ทิ้งงาน รับประกันงานซ่อม 30 วัน"
    },
    "🔨 ช่างรับเหมา": {
        "name": "ช่างสมชาย รับเหมาต่อเติม",
        "product": "ต่อเติมครัว โรงจอดรถ ปูกระเบื้อง ซ่อมแซมทั่วไป",
        "customer": "คนในชุมชนระแวกใกล้เคียง 10 กม., ผู้สูงอายุ",
        "usp": "เป็นคนในพื้นที่ ไว้ใจได้ ปรึกษาหน้างานฟรี"
    },
    "🏍️ ซ่อมมอไซค์": {
        "name": "อู่ช่างบอย มอไซค์ซิ่ง",
        "product": "ซ่อมมอเตอร์ไซค์ ถ่ายน้ำมันเครื่อง ปะยาง แต่งรถ",
        "customer": "วินมอเตอร์ไซค์, นักเรียน, คนทำงานโรงงาน",
        "usp": "เปิดเช้าปิดดึก มีรถกระบะไปรับรถเสียถึงที่"
    },
    "🥬 ผักไฮโดรฯ": {
        "name": "กรีนฟาร์ม ไฮโดรโปนิกส์",
        "product": "ผักสลัด (กรีนโอ๊ค, เรดโอ๊ค) ปลอดสารพิษ สดใหม่",
        "customer": "คนรักสุขภาพ, ร้านสเต็ก, ร้านสลัดโรล",
        "usp": "ผักสดตัดใหม่ทุกเช้า ไม่ใช้ยาฆ่าแมลง มี QR ตรวจสอบ"
    },
    "☕ ร้านกาแฟ": {
        "name": "กาแฟบ้านทุ่ง",
        "product": "กาแฟสด เมนูน้ำชง โกโก้ ชาเขียว ขนมปังปิ้ง",
        "customer": "คนในชุมชน, เกษตรกรพักเที่ยง, ขาจรขับรถผ่าน",
        "usp": "ราคาเข้าถึงง่าย (25-40 บาท) รสชาติเข้มข้น บรรยากาศกันเอง"
    }
}
EXAMPLE_OPTIONS = ["✨ เริ่มต้นใหม่ (ล้างข้อมูล)"] + list(EXAMPLES_DATA.keys())

# --- ฟังก์ชันสร้างไฟล์ Word (แก้ไขจุด Bug แล้ว) ---
def create_word_docx(data, mode, business_name):
    doc = Document()
    
    # Style พื้นฐาน
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Sarabun' 
    font.size = Pt(11)

    # Header
    header = doc.add_heading(level=1)
    run = header.add_run(f"{mode}: {business_name}")
    run.font.name = 'Sarabun'
    run.font.color.rgb = RGBColor(0x4a, 0x14, 0x8c) 
    
    doc.add_paragraph("สร้างโดย AI Business Planner - กรมพัฒนาฝีมือแรงงาน").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("-" * 50)

    if mode == "BMC":
        # สร้างตาราง 3x5 สำหรับ BMC
        table = doc.add_table(rows=3, cols=5)
        table.style = 'Table Grid'
        
        # Row 1 (ใช้ table.cell ได้ปกติ)
        table.cell(0, 0).text = f"KP: Key Partners\n\n{data.get('key_partners', '-')}"
        table.cell(0, 1).text = f"KA: Key Activities\n\n{data.get('key_activities', '-')}"
        table.cell(0, 2).text = f"VP: Value Propositions\n\n{data.get('value_propositions', '-')}"
        table.cell(0, 3).text = f"CR: Customer Relationships\n\n{data.get('customer_relationships', '-')}"
        table.cell(0, 4).text = f"CS: Customer Segments\n\n{data.get('customer_segments', '-')}"
        
        # Row 2 (Merge)
        table.cell(1, 0).merge(table.cell(2, 0))
        table.cell(1, 1).text = f"KR: Key Resources\n\n{data.get('key_resources', '-')}"
        table.cell(1, 2).merge(table.cell(2, 2))
        table.cell(1, 3).text = f"CH: Channels\n\n{data.get('channels', '-')}"
        table.cell(1, 4).merge(table.cell(2, 4))
        
        # Row 3 (Merge empty spaces)
        table.cell(2, 1).merge(table.cell(2, 1))
        table.cell(2, 3).merge(table.cell(2, 3))
        
        # --- จุดที่แก้ไข (ใช้ .cells แทน .cell) ---
        row_cr = table.add_row()
        row_cr.cells[0].merge(row_cr.cells[1]) # แก้ไขจาก .cell(0) เป็น .cells[0]
        row_cr.cells[0].text = f"C$: Cost Structure\n\n{data.get('cost_structure', '-')}"
        
        row_cr.cells[2].merge(row_cr.cells[4]) # แก้ไขจาก .cell(2) เป็น .cells[2]
        row_cr.cells[2].text = f"R$: Revenue Streams\n\n{data.get('revenue_streams', '-')}"

    elif mode == "VPC":
        # สร้างตาราง 3 Row x 2 Col สำหรับ VPC แบบจับคู่
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        # หัวตาราง (เข้าถึงผ่าน .rows[].cells[])
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "📦 Value Map (ฝั่งสินค้า)"
        hdr_cells[1].text = "👤 Customer Profile (ฝั่งลูกค้า)"
        
        # Row 1: Products <-> Jobs
        row1 = table.add_row().cells
        row1[0].text = f"Products & Services:\n\n{data.get('products_services', '-')}"
        row1[1].text = f"Customer Jobs:\n\n{data.get('customer_jobs', '-')}"
        
        # Row 2: Pain Relievers <-> Pains
        row2 = table.add_row().cells
        row2[0].text = f"Pain Relievers:\n\n{data.get('pain_relievers', '-')}"
        row2[1].text = f"Pains:\n\n{data.get('pains', '-')}"
        
        # Row 3: Gain Creators <-> Gains
        row3 = table.add_row().cells
        row3[0].text = f"Gain Creators:\n\n{data.get('gain_creators', '-')}"
        row3[1].text = f"Gains:\n\n{data.get('gains', '-')}"

    # บันทึก
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- CSS รวมญาติ ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Sarabun:wght@300;400;600&display=swap');
    html, body, [class*="css"] { font-family: 'Sarabun', sans-serif; }

    /* Header */
    .header-container {
        background: linear-gradient(135deg, #4a148c 0%, #7b1fa2 100%);
        padding: 25px; border-radius: 12px; color: white; margin-bottom: 25px;
        box-shadow: 0 4px 8px rgba(0,0,0,0.15); display: flex; align-items: center; gap: 20px;
    }
    .logo-img { width: 90px; height: 90px; object-fit: contain; background-color: white; border-radius: 50%; padding: 5px; border: 3px solid #FFC107; flex-shrink: 0; }
    .header-text { text-align: left; width: 100%; }
    .header-main { font-size: 1.8rem; font-weight: bold; margin: 0; color: #FFF; line-height: 1.2; }
    .header-sub { font-size: 1.1rem; font-weight: 400; margin-bottom: 5px; color: #FFD54F; }
    .header-line { border-bottom: 3px solid #FFC107; width: 80px; margin: 10px 0; }
    
    @media only screen and (max-width: 768px) {
        .header-container { flex-direction: column; text-align: center; }
        .logo-img { width: 70px; height: 70px; }
    }

    /* Radio Buttons */
    .stRadio [role=radiogroup] { gap: 10px; flex-wrap: wrap; justify-content: center; }
    .stRadio label[data-baseweb="radio"] {
        background-color: #ffffff; border: 2px solid #9c27b0; color: #6a1b9a;
        padding: 8px 16px; border-radius: 20px; cursor: pointer; transition: all 0.2s;
        display: flex; align-items: center; justify-content: center; min-width: 110px; margin: 0 !important;
    }
    .stRadio label[data-baseweb="radio"]:hover { background-color: #f3e5f5; transform: translateY(-2px); }
    .stRadio label[data-baseweb="radio"][aria-checked="true"] {
        background-color: #4a148c !important; color: #ffffff !important; border-color: #4a148c !important;
        box-shadow: 0 4px 10px rgba(74, 20, 140, 0.4); font-weight: bold;
    }
    .stRadio label[data-baseweb="radio"] > div:first-child { display: none; }

    /* Box Style (Shared) */
    .box { background-color: #ffffff; border: 1px solid #e0e0e0; border-radius: 8px; padding: 15px; color: #333; box-shadow: 0 2px 4px rgba(0,0,0,0.05); height: 100%; min-height: 160px; }
    .box h4 { margin-top: 0; color: #4a148c; font-size: 1rem; font-weight: bold; margin-bottom: 10px; border-bottom: 2px solid #f3e5f5; padding-bottom: 5px; min-height: 55px; }
    .box p { font-size: 0.9rem; line-height: 1.6; white-space: pre-wrap; color: #555; margin: 0; }

    /* --- BMC CSS --- */
    .bmc-grid { display: grid; grid-template-columns: repeat(5, 1fr); grid-template-rows: repeat(3, minmax(180px, auto)); gap: 12px; margin-top: 20px; }
    .kp { grid-area: 1 / 1 / 3 / 2; background-color: #f3e5f5; }
    .ka { grid-area: 1 / 2 / 2 / 3; }
    .kr { grid-area: 2 / 2 / 3 / 3; }
    .vp { grid-area: 1 / 3 / 3 / 4; background-color: #fffde7; border: 2px solid #FFC107; }
    .cr { grid-area: 1 / 4 / 2 / 5; }
    .ch { grid-area: 2 / 4 / 3 / 5; }
    .cs { grid-area: 1 / 5 / 3 / 6; background-color: #f3e5f5; }
    .co { grid-area: 3 / 1 / 4 / 3; background-color: #fff5f5; border: 1px dashed #dc3545; }
    .rs { grid-area: 3 / 3 / 4 / 6; background-color: #f0fff4; border: 1px dashed #28a745; }

    /* --- VPC CSS (New Layout) --- */
    .vpc-section { margin-top: 20px; display: flex; flex-direction: column; gap: 20px; }
    .vpc-pair-row { display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }
    .vpc-box-left { background-color: #e8eaf6; border: 2px dashed #3f51b5; padding: 20px; border-radius: 12px; }
    .vpc-box-right { background-color: #e0f2f1; border: 2px dashed #009688; padding: 20px; border-radius: 12px; }
    
    @media only screen and (max-width: 768px) {
        .bmc-grid, .vpc-pair-row { display: flex; flex-direction: column; }
    }

    /* Buttons */
    button[kind="primary"] { background-color: #4a148c !important; border: none; color: white !important; border-radius: 8px; font-weight: bold; }
    button:hover { transform: scale(1.02); transition: 0.2s; }

    /* --- Footer CSS --- */
    .footer-container {
        background: linear-gradient(135deg, #4a148c 0%, #7b1fa2 100%);
        padding: 20px; border-radius: 12px; color: white; text-align: center;
        margin-top: 40px; box-shadow: 0 -2px 5px rgba(0,0,0,0.1);
    }
    .footer-credit { font-weight: bold; color: #FFD54F; }
    .footer-line { border-bottom: 2px solid #FFC107; width: 50px; margin: 10px auto; }
</style>
""", unsafe_allow_html=True)

# --- Header ---
st.markdown(f"""
<div class="header-container">
    <img src="{logo_src}" class="logo-img" alt="DSD Logo">
    <div class="header-text">
        <div class="header-sub">กรมพัฒนาฝีมือแรงงาน</div>
        <div class="header-main">สำนักงานพัฒนาฝีมือแรงงานสกลนคร</div>
        <div class="header-line"></div>
        <div class="header-desc">ระบบสร้างแผนธุรกิจอัตโนมัติ (AI Business Planner)</div>
    </div>
</div>
""", unsafe_allow_html=True)

# --- AI Function ---
def generate_business_plan(mode, business, product, customer, strength):
    if not api_key: st.error("ไม่พบ API Key"); return None
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.5-flash-lite')
    
    if mode == "BMC":
        task = "Create Business Model Canvas (BMC)"
        keys = "key_partners, key_activities, key_resources, value_propositions, customer_relationships, channels, customer_segments, cost_structure, revenue_streams"
    else: # VPC
        task = "Create Value Proposition Canvas (VPC)"
        keys = "products_services, pain_relievers, gain_creators, customer_jobs, pains, gains"

    prompt = f"""
    Role: Business Consultant.
    Task: {task} for Business: "{business}" | Product: "{product}" | Customer: "{customer}" | USP: "{strength}"
    Constraint: Return ONLY JSON object. Thai Language. Use bullet points (-) for lists.
    Required JSON Keys: {keys}
    """
    try:
        response = model.generate_content(prompt, generation_config={"response_mime_type": "application/json"})
        text_res = response.text.strip()
        if "```json" in text_res: text_res = text_res.replace("```json", "").replace("```", "")
        match = re.search(r'\{[\s\S]*\}', text_res)
        data = json.loads(match.group(0) if match else text_res)
        cleaned = {}
        for k, v in data.items():
            if isinstance(v, list): v = "\n".join([f"- {str(i)}" for i in v])
            elif isinstance(v, dict): v = "\n".join([f"- {item}" for item in v.values()])
            else: v = str(v)
            cleaned[k] = v.replace("['", "").replace("']", "").replace('["', '').replace('"]', '')
        return cleaned
    except Exception as e:
        st.error(f"AI Error: {e}")
        return {}

# --- Logic ---
if 'form_data' not in st.session_state:
    st.session_state['form_data'] = {'name': '', 'product': '', 'customer': '', 'usp': ''}

def update_form():
    sel = st.session_state['radio_select']
    st.session_state['form_data'] = EXAMPLES_DATA.get(sel, {'name': '', 'product': '', 'customer': '', 'usp': ''}).copy()

st.write("##### 💡 เลือกตัวอย่างธุรกิจ:")
curr_name = st.session_state['form_data']['name']
curr_idx = 0
if curr_name:
    for i, k in enumerate(EXAMPLES_DATA):
        if EXAMPLES_DATA[k]['name'] == curr_name: curr_idx = i + 1; break

st.radio("ตัวเลือก", options=EXAMPLE_OPTIONS, index=curr_idx, horizontal=True, label_visibility="collapsed", key="radio_select", on_change=update_form)
st.divider()

# --- Main Form ---
with st.form("biz_form"):
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**1. ชื่อธุรกิจ**"); b_name = st.text_input("ชื่อ", st.session_state['form_data']['name'], label_visibility="collapsed")
        st.markdown("**3. ลูกค้า**"); b_cust = st.text_input("ลูกค้า", st.session_state['form_data']['customer'], label_visibility="collapsed")
    with c2:
        st.markdown("**2. สินค้า/บริการ**"); b_prod = st.text_area("สินค้า", st.session_state['form_data']['product'], label_visibility="collapsed", height=104)
        st.markdown("**4. จุดเด่น**"); b_usp = st.text_input("จุดเด่น", st.session_state['form_data']['usp'], label_visibility="collapsed")
    
    st.write("---")
    st.write("**เลือกรูปแบบผลลัพธ์:**")
    b_col1, b_col2 = st.columns(2)
    with b_col1: submit_bmc = st.form_submit_button("🚀 สร้างโมเดลธุรกิจ (BMC)", type="primary", use_container_width=True)
    with b_col2: submit_vpc = st.form_submit_button("🎯 สร้างแผนคุณค่า (VPC)", type="primary", use_container_width=True)

# --- Display Output ---
if submit_bmc or submit_vpc:
    if not b_name:
        st.warning("⚠️ กรุณากรอกข้อมูลชื่อธุรกิจก่อนครับ")
    else:
        mode = "BMC" if submit_bmc else "VPC"
        with st.spinner(f"⏳ กำลังวิเคราะห์ข้อมูลและสร้าง {mode}..."):
            data = generate_business_plan(mode, b_name, b_prod, b_cust, b_usp)
            
            if data:
                if mode == "BMC":
                    st.success("📊 **Business Model Canvas (BMC)**")
                    bmc_html = f"""
                    <div class="bmc-grid">
                        <div class="box kp"><h4>🤝 Key Partners<br>(พันธมิตรหลัก)</h4><p>{data.get('key_partners', '-')}</p></div>
                        <div class="box ka"><h4>⚙️ Key Activities<br>(กิจกรรมหลัก)</h4><p>{data.get('key_activities', '-')}</p></div>
                        <div class="box kr"><h4>🧱 Key Resources<br>(ทรัพยากรหลัก)</h4><p>{data.get('key_resources', '-')}</p></div>
                        <div class="box vp"><h4>🎁 Value Propositions<br>(คุณค่าที่ส่งมอบ)</h4><p>{data.get('value_propositions', '-')}</p></div>
                        <div class="box cr"><h4>❤️ Customer Relationships<br>(ความสัมพันธ์ลูกค้า)</h4><p>{data.get('customer_relationships', '-')}</p></div>
                        <div class="box ch"><h4>🚚 Channels<br>(ช่องทางเข้าถึง)</h4><p>{data.get('channels', '-')}</p></div>
                        <div class="box cs"><h4>👥 Customer Segments<br>(กลุ่มลูกค้าหลัก)</h4><p>{data.get('customer_segments', '-')}</p></div>
                        <div class="box co"><h4>💰 Cost Structure<br>(โครงสร้างต้นทุน)</h4><p>{data.get('cost_structure', '-')}</p></div>
                        <div class="box rs"><h4>💵 Revenue Streams<br>(กระแสรายได้)</h4><p>{data.get('revenue_streams', '-')}</p></div>
                    </div>
                    """
                    st.markdown(bmc_html, unsafe_allow_html=True)
                
                elif mode == "VPC":
                    st.success("🎯 **Value Proposition Canvas (VPC) - แบบจับคู่**")
                    vpc_html = f"""
                    <div class="vpc-section">
                        <div class="vpc-pair-row">
                            <div class="vpc-box-left">
                                <div style="color:#3f51b5; font-weight:bold; margin-bottom:10px;">📦 Value Map (ฝั่งสินค้า)</div>
                                <div class="box"><h4>🛍️ Products & Services<br>(สินค้าและบริการ)</h4><p>{data.get('products_services', '-')}</p></div>
                            </div>
                            <div class="vpc-box-right">
                                <div style="color:#00695c; font-weight:bold; margin-bottom:10px;">👤 Customer Profile (ฝั่งลูกค้า)</div>
                                <div class="box"><h4>📝 Customer Jobs<br>(งานที่ลูกค้าต้องทำ)</h4><p>{data.get('customer_jobs', '-')}</p></div>
                            </div>
                        </div>
                        <div class="vpc-pair-row">
                            <div class="vpc-box-left">
                                <div class="box"><h4>💊 Pain Relievers<br>(ตัวช่วยแก้ปัญหา)</h4><p>{data.get('pain_relievers', '-')}</p></div>
                            </div>
                            <div class="vpc-box-right">
                                <div class="box"><h4>😫 Pains<br>(ปัญหา/ความยุ่งยาก)</h4><p>{data.get('pains', '-')}</p></div>
                            </div>
                        </div>
                        <div class="vpc-pair-row">
                            <div class="vpc-box-left">
                                <div class="box"><h4>⚡ Gain Creators<br>(สิ่งที่ช่วยสร้างประโยชน์)</h4><p>{data.get('gain_creators', '-')}</p></div>
                            </div>
                            <div class="vpc-box-right">
                                <div class="box"><h4>😍 Gains<br>(ประโยชน์ที่คาดหวัง)</h4><p>{data.get('gains', '-')}</p></div>
                            </div>
                        </div>
                    </div>
                    """
                    st.markdown(vpc_html, unsafe_allow_html=True)

                # --- Export Button ---
                st.write("---")
                st.write(f"📥 **ดาวน์โหลดผลลัพธ์ ({mode}):**")
                docx_file = create_word_docx(data, mode, b_name)
                st.download_button(
                    label=f"📄 ดาวน์โหลดเป็น Word (.docx)",
                    data=docx_file,
                    file_name=f"{mode}_{b_name.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

# --- Footer ---
st.markdown("""
<div class="footer-container">
    <div class="footer-credit"> |-- เทอดศิลป์ โสมูล --|</div>
    <div class="footer-line"></div>
    <p>© 2025 สำนักงานพัฒนาฝีมือแรงงานสกลนคร | AI Powered by Gemini</p>
</div>
""", unsafe_allow_html=True)